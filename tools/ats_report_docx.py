from pathlib import Path

from docx import Document


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        if item:
            doc.add_paragraph(item, style="List Bullet")


def create_ats_report_docx(ats_review, output_path: str | Path) -> Path:
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    doc.add_heading("ATS Optimization Report", level=1)

    doc.add_heading("Scores", level=2)
    doc.add_paragraph(f"ATS Score: {ats_review.ats_score}")
    doc.add_paragraph(f"Keyword Match Score: {ats_review.keyword_match_score}")

    doc.add_heading("Matched Keywords", level=2)
    add_bullets(doc, ats_review.matched_keywords)

    doc.add_heading("Missing Keywords", level=2)
    add_bullets(doc, ats_review.missing_keywords)

    doc.add_heading("Resume Strengths", level=2)
    add_bullets(doc, ats_review.resume_strengths)

    doc.add_heading("Resume Weaknesses", level=2)
    add_bullets(doc, ats_review.resume_weaknesses)

    doc.add_heading("Improvement Suggestions", level=2)
    add_bullets(doc, ats_review.improvement_suggestions)

    doc.add_heading("Regeneration Recommendation", level=2)
    doc.add_paragraph(
        "Regenerate Resume: "
        + ("Yes" if ats_review.should_regenerate_resume else "No")
    )

    doc.add_heading("Final Recommendation", level=2)
    doc.add_paragraph(ats_review.final_recommendation)

    doc.save(output_path)
    return output_path