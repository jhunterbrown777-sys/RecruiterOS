from pathlib import Path

from docx import Document


def add_bullet_list(doc: Document, items: list[str]) -> None:
    for item in items:
        if item:
            doc.add_paragraph(item, style="List Bullet")


def create_resume_docx(resume, output_path: str | Path) -> Path:
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    doc.add_heading(resume.resume_title, level=1)

    doc.add_heading("Professional Summary", level=2)
    doc.add_paragraph(resume.professional_summary)

    doc.add_heading("Core Skills", level=2)
    add_bullet_list(doc, resume.core_skills)

    doc.add_heading("Technical Skills", level=2)
    add_bullet_list(doc, resume.technical_skills)

    doc.add_heading("Professional Experience", level=2)
    add_bullet_list(doc, resume.selected_experience)

    doc.add_heading("Projects", level=2)
    add_bullet_list(doc, resume.selected_projects)

    doc.add_heading("Education", level=2)
    add_bullet_list(doc, resume.education)

    doc.add_heading("Certifications & Training", level=2)
    add_bullet_list(doc, resume.certifications)

    doc.add_heading("ATS Keywords Included", level=2)
    add_bullet_list(doc, resume.ats_keywords_included)

    doc.add_heading("Resume Strategy", level=2)
    doc.add_paragraph(resume.resume_strategy)

    if resume.honesty_notes:
        doc.add_heading("Honesty Notes", level=2)
        add_bullet_list(doc, resume.honesty_notes)

    doc.save(output_path)
    return output_path