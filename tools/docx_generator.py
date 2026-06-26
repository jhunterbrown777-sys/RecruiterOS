from pathlib import Path
from docx import Document


def create_docx(title: str, content: str, output_path: str):
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    doc.add_heading(title, level=1)

    for paragraph in content.split("\n"):
        paragraph = paragraph.strip()
        if paragraph:
            doc.add_paragraph(paragraph)

    doc.save(output_path)
    return output_path